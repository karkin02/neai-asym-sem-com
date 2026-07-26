from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "FAULTY_COMMAND_INCIDENT_REPORT.md"
TXT_OUT = ROOT / "FAULTY_COMMAND_INCIDENT_REPORT.txt"
DOCX_OUT = ROOT / "FAULTY_COMMAND_INCIDENT_REPORT.docx"

INK = RGBColor(31, 41, 55)
ACCENT = RGBColor(31, 78, 121)
MUTED = RGBColor(92, 103, 115)
CODE_FILL = "F3F4F6"


def set_font(run, name="Aptos", size=10.5, bold=False, color=INK):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color


def shade_paragraph(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end])
    set_font(run, size=9, color=MUTED)


def plain_text(markdown):
    lines = []
    in_code = False
    for raw in markdown.splitlines():
        line = raw.rstrip()
        if line.startswith("```"):
            in_code = not in_code
            continue
        if not in_code:
            line = re.sub(r"^#{1,6}\s+", "", line)
            line = re.sub(r"^- ", "* ", line)
            line = line.replace("`", "")
        lines.append(line)
    return "\n".join(lines).strip() + "\n"


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.78)
    section.left_margin = Inches(0.92)
    section.right_margin = Inches(0.92)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12

    for style_name, size, before, after in (
        ("Heading 1", 15, 14, 5),
        ("Heading 2", 12, 10, 4),
    ):
        style = doc.styles[style_name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = ACCENT
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.text = "INCIDENT REPORT  |  CODEX-GENERATED PACKAGE UNINSTALL"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_font(header.runs[0], size=8.5, bold=True, color=MUTED)

    footer = section.footer.paragraphs[0]
    add_page_number(footer)


def add_title_block(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("CODEX FAULTY COMMAND")
    set_font(run, name="Aptos Display", size=22, bold=True, color=INK)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("Incident report and request for credit review")
    set_font(run, name="Aptos Display", size=13, color=MUTED)

    metadata = [
        ("Incident date", "July 23, 2026"),
        ("Environment", "Windows 25H2; Python 3.11.9"),
        ("Affected scope", "System-wide Python 3.11 package environment"),
        ("Status", "Recovery in progress"),
    ]
    for label, value in metadata:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1.5)
        set_font(p.add_run(f"{label}: "), size=9.5, bold=True)
        set_font(p.add_run(value), size=9.5)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    p_pr = p._p.get_or_add_pPr()
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "10")
    bottom.set(qn("w:color"), "1F4E79")
    border.append(bottom)
    p_pr.append(border)


def build_docx(markdown):
    doc = Document()
    configure_document(doc)
    add_title_block(doc)

    lines = markdown.splitlines()
    in_code = False
    code_lines = []
    first_heading = True

    for raw in lines:
        line = raw.rstrip()
        if line.startswith("```"):
            if in_code:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.18)
                p.paragraph_format.right_indent = Inches(0.18)
                p.paragraph_format.space_before = Pt(3)
                p.paragraph_format.space_after = Pt(7)
                shade_paragraph(p, CODE_FILL)
                run = p.add_run("\n".join(code_lines))
                set_font(run, name="Consolas", size=8.2, color=INK)
                code_lines = []
            in_code = not in_code
            continue
        if in_code:
            code_lines.append(line)
            continue
        if line.startswith("# "):
            if first_heading:
                first_heading = False
            continue
        if line.startswith("## "):
            doc.add_paragraph(line[3:], style="Heading 1")
            continue
        if line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Inches(0.32)
            p.paragraph_format.first_line_indent = Inches(-0.16)
            p.paragraph_format.space_after = Pt(3)
            p.add_run(line[2:].replace("`", ""))
            continue
        if re.match(r"^\d+\.\s", line):
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.left_indent = Inches(0.35)
            p.paragraph_format.first_line_indent = Inches(-0.18)
            p.paragraph_format.space_after = Pt(3)
            p.add_run(re.sub(r"^\d+\.\s", "", line).replace("`", ""))
            continue
        if not line:
            continue

        p = doc.add_paragraph()
        parts = re.split(r"(`[^`]+`)", line)
        for part in parts:
            if part.startswith("`") and part.endswith("`"):
                run = p.add_run(part[1:-1])
                set_font(run, name="Consolas", size=9, color=INK)
            else:
                p.add_run(part)

    doc.save(DOCX_OUT)


def main():
    markdown = SOURCE.read_text(encoding="utf-8")
    TXT_OUT.write_text(plain_text(markdown), encoding="utf-8")
    build_docx(markdown)
    print(TXT_OUT)
    print(DOCX_OUT)


if __name__ == "__main__":
    main()
